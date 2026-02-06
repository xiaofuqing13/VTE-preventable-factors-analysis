# -*- coding: utf-8 -*-
"""
11_update_report_r.py
在报告中补充R语言6模型完整结果和Python/R全面对比
替换原有的10.8节（仅RF+XGBoost验证），扩展为完整的R语言6模型分析
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'C:\Users\Administrator\Desktop\2.5.388'

def set_cell_font(cell, text, size=10, bold=False, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading_black(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table_from_df(doc, df, caption=None):
    if caption:
        add_para(doc, caption, size=10, bold=True)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        set_cell_font(table.rows[0].cells[j], col, size=9, bold=True)
    for i in range(len(df)):
        for j, col in enumerate(df.columns):
            val = df.iloc[i, j]
            if isinstance(val, float):
                val = f'{val:.4f}' if abs(val) < 10 else f'{val:.2f}'
            set_cell_font(table.rows[i + 1].cells[j], str(val), size=9)
    return table

# 打开现有报告
doc = Document(os.path.join(BASE, 'VTE影响因素分析报告.docx'))

# 在末尾追加R语言完整6模型分析
add_heading_black(doc, '10.8 R语言6种模型完整分析', level=2)
add_para(doc, '使用R语言（R 4.5.2）独立实现全部6种机器学习模型，验证Python分析结果的可重复性。R语言使用的包：randomForest（Random Forest）、e1071（SVM和Naive Bayes）、xgboost（XGBoost）、rpart（Decision Tree）、class（KNN）、pROC（ROC分析）。')

# R模型参数
r_params = pd.read_csv(os.path.join(BASE, 'r_ml_model_params.csv'))
add_table_from_df(doc, r_params, caption='表27 R语言模型参数设置')

# R模型测试集对比
r_comp = pd.read_csv(os.path.join(BASE, 'r_ml_model_comparison.csv'))
add_table_from_df(doc, r_comp, caption='表28 R语言模型测试集性能对比（按AUC排序）')

add_para(doc, 'R语言模型性能排序为：XGBoost(0.982) > Decision Tree(0.913) > Random Forest(0.912) > SVM(0.808) > KNN(0.505) > Naive Bayes(0.500)，与Python结果基本一致，XGBoost均为最佳模型。')

# R训练集vs测试集
r_tt = pd.read_csv(os.path.join(BASE, 'r_ml_train_test_comparison.csv'))
add_table_from_df(doc, r_tt, caption='表29 R语言各模型训练集与测试集性能对比')

# R外部验证
r_ext = pd.read_csv(os.path.join(BASE, 'r_ml_external_validation.csv'))
add_table_from_df(doc, r_ext, caption='表30 R语言模型外部验证性能')

# Python vs R 全对比
add_heading_black(doc, '10.9 Python与R语言全面对比', level=2)
add_para(doc, '对6种模型在测试集和外部验证集上的Python与R语言AUC进行全面对比，评估结果一致性。')

cross = pd.read_csv(os.path.join(BASE, 'r_vs_python_full_comparison.csv'))
add_table_from_df(doc, cross, caption='表31 Python与R语言6种模型AUC全面对比')

add_para(doc, 'Python与R全面对比结果：')
add_para(doc, '（1）XGBoost测试集AUC完全一致（均为0.982），外部验证差异仅0.006，一致性最好。')
add_para(doc, '（2）Random Forest测试集差异0.016，SVM差异0.006，Decision Tree差异0.068，均在可接受范围内。')
add_para(doc, '（3）KNN两种语言测试集AUC均约0.5，一致反映了高维空间中KNN的局限性。')
add_para(doc, '（4）Naive Bayes测试集差异0.175，为最大差异项，原因是Python sklearn.GaussianNB与R e1071::naiveBayes在概率估计细节上的实现差异，但两者性能均较差。')
add_para(doc, '（5）总体结论：12组对比中11组差异<0.1，核心模型（XGBoost/RF/SVM）高度一致，验证了分析结果的可靠性和可重复性。')

# 更新10.10小结
add_heading_black(doc, '10.10 小结', level=2)
add_para(doc, '本章使用Python和R两种语言、6种机器学习模型对潜在可预防VTE进行风险预测分析，主要发现：')
add_para(doc, '（1）XGBoost为最佳预测模型：Python和R测试集AUC均为0.982，F1=0.884，综合判别性能显著优于其他模型。')
add_para(doc, '（2）SHAP可解释性分析揭示了关键预测特征，与传统单因素/多因素分析结果高度一致：90天前就诊史（SHAP=1.95, 多因素OR=38.24）、机械预防（SHAP=1.04, OR=0.22）、出血风险评估（SHAP=0.66, OR=0.30）是最重要的因素。')
add_para(doc, '（3）外部验证（3-31后88例）显示模型具有一定泛化能力（最佳AUC=0.743），精确率高（>0.9）但敏感性偏低。')
add_para(doc, '（4）Python与R交叉验证证实了分析结果的可重复性（12组对比中11组差异<0.1）。')
add_para(doc, '（5）综合统计学方法和机器学习方法，90天前就诊史、机械预防、出血风险评估是潜在可预防VTE最稳健的影响因素，建议临床重点关注。')

# 保存
doc.save(os.path.join(BASE, 'VTE影响因素分析报告.docx'))
print('报告已更新：补充R语言6模型完整结果')
print('  10.8 R语言6种模型完整分析 (表27-30)')
print('  10.9 Python与R全面对比 (表31)')
print('  10.10 小结')
