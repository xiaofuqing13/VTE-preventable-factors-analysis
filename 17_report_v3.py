# -*- coding: utf-8 -*-
"""
17_report_v3.py - 严格按照模板(数据分析报告模板(1).docx)的全部元素生成报告
模板要素对照：
  一、研究概述【无需填写→简述】
  二、变量筛选
    2.1 单因素(4类表) → 表1-4
    2.2 多因素
      2.2.1 自变量赋值 → 表5
      2.2.2 多因素回归 → 表6
      2.2.3 模型构建(每模型:混淆矩阵图+评价表+ROC图) → 表7-12, 图1-12
      2.2.4 模型评价 → 表13
      2.2.5 SHAP(特征重要性图+概要图+交互热力图) → 图13-15
  三、3-31后分析
    3.1 单因素(4类) → 表14-17
    3.2 多因素 → 表18-19
    3.3 结局指标
      3.3.1 前后基线对比 → 表20
      3.3.2 HA-VTE发生率 → 表21
      3.3.3 潜在可预防VTE → 表22
      3.3.4 规范预防 → 表23-26
  四、结论
  五、方法一致性验证 → 表27
"""
import os, re, math
import pandas as pd
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'C:\Users\Administrator\Desktop\2.5.388'
doc = Document()

# ===== 默认样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'; font.size = Pt(11); font.color.rgb = RGBColor(0,0,0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 工具函数 =====
def sr(run, size=11, bold=False):
    run.font.size = Pt(size); run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0,0,0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold

def at(text, size=11, bold=False, align=None):
    p = doc.add_paragraph(); run = p.add_run(text); sr(run, size, bold)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def title(t):    return at(t, 18, True, 'center')
def ch(t):       return at(t, 16, True)
def sec(t):      return at(t, 14, True)
def sub(t):      return at(t, 11, True)
def body(t):     return at(t, 11, False)
def cap(t):      return at(t, 11, True)

def sc(cell, text, size=9, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]; run = p.add_run(str(text)); sr(run, size, bold)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def mt(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers): sc(table.rows[0].cells[j], h, 9, True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row): sc(table.rows[i+1].cells[j], fv(val), 9)
    return table

def fv(v):
    if isinstance(v, float):
        if pd.isna(v): return ''
        if abs(v) < 0.0001: return f'{v:.6f}'
        if abs(v) < 0.001: return f'{v:.4f}'
        if abs(v) < 10: return f'{v:.4f}'
        return f'{v:.2f}'
    return str(v)

def img(path, caption=None, w=5.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption: at(caption, 9, False, 'center')

def rc(name): return pd.read_csv(os.path.join(BASE, name))

# ===== 变量分类 =====
def clf(name):
    n = str(name)
    bio = ['嗜酸','D-二聚体','DD','APTT','凝血','钾离子','纤维蛋白原','白蛋白','血红蛋白',
           '总钙','白细胞','中性粒','钠','血小板','淋巴','单核','嗜碱','红细胞','血糖',
           'Fbg','FEU','尿素','肌酐','总蛋白','球蛋白','ALT','AST','胆红素','PT','INR','TT','钙','磷']
    for k in bio:
        if k in n: return '生化'
    treat = ['机械预防','药物预防','预防措施','出血风险评估','规范预防',
             'VTE首次中高危评分日期与机械预防','VTE首次中高危评分日期与药物预防',
             '首次VTE中高风险评分日期与机械预防','首次DD与入院','首次异常D','首次异常DD']
    for k in treat:
        if k in n: return '治疗'
    dis = ['医院相关性VTE','我院相关VTE','90天前','入院前90天','主要疾病诊断','呼吸系统',
           '糖尿病','高血压','肿瘤','炎症','感染','VTE中高危评分','抗凝禁忌','COPD',
           '既往','肺栓塞','脑梗','脑内出血','心衰','肺炎','冠心病','肝硬化',
           'VTE评估量表','VTE诊断日期','潜在可预防']
    for k in dis:
        if k in n: return '疾病'
    return '一般'

# ===== 计算辅助 =====
def parse_cat(s):
    m = re.match(r'(\d+)/(\d+)', str(s))
    return (int(m.group(1)), int(m.group(2))) if m else (None, None)

def parse_cont(s):
    m = re.match(r'([\d.]+)±([\d.]+)', str(s))
    return (float(m.group(1)), float(m.group(2))) if m else (None, None)

def total_stat(row, nn, np_):
    vt = row['类型']
    nc = [c for c in row.index if '阴性' in c][0]
    pc = [c for c in row.index if '阳性' in c][0]
    if vt == '分类':
        a, _ = parse_cat(row[pc]); c, _ = parse_cat(row[nc])
        if a is not None and c is not None:
            tn = nn + np_; ty = a + c
            ts = f'{ty}/{tn} ({ty/tn*100:.1f}%)'
            b = np_ - a; d = nn - c
            try:
                tab = np.array([[a,b],[c,d]])
                chi2, _, _, _ = stats.chi2_contingency(tab, correction=False)
                return ts, f'{chi2:.3f}'
            except: return ts, '-'
        return '-', '-'
    elif vt == '连续':
        m1, s1 = parse_cont(row[pc]); m2, s2 = parse_cont(row[nc])
        if m1 is not None:
            tn = nn + np_; pm = (np_*m1+nn*m2)/tn
            pv = ((np_-1)*s1**2+(nn-1)*s2**2)/(tn-2)
            ps = math.sqrt(pv) if pv > 0 else 0
            se = math.sqrt(s1**2/np_+s2**2/nn)
            tv = (m1-m2)/se if se > 0 else 0
            return f'{pm:.2f}±{ps:.2f}', f'{tv:.3f}'
        return '-', '-'
    return '-', '-'

def sf_table(bdf, nn, np_, cat, raw_data=None):
    """单因素表格，增加均值±标准差列（批注A/B/E/K）"""
    cdf = bdf[bdf['变量'].apply(clf) == cat]
    sig = cdf[cdf['P值'] < 0.05]
    if len(sig) == 0: return [], 0, len(cdf)
    nc = [c for c in sig.columns if '阴性' in c][0]
    pc = [c for c in sig.columns if '阳性' in c][0]
    # 生化关键词
    _bio = ['嗜酸','D-二聚体','DD','APTT','凝血','钾离子','纤维蛋白原','白蛋白','血红蛋白',
            '总钙','白细胞','中性粒','红细胞','Fbg','FEU','PT','MCH','MCHC','Braden','钠','钾']
    rows = []
    for _, r in sig.iterrows():
        ts, st = total_stat(r, nn, np_)
        vname = r['变量']
        # 计算均值±标准差和检测例数
        mean_sd = '-'
        if raw_data is not None and r['类型'] == '连续' and vname in raw_data.columns:
            vals = raw_data[vname].dropna()
            is_bio = any(k in vname for k in _bio)
            if is_bio:
                vals = vals[vals != 0]
            if len(vals) > 0:
                mean_sd = f'{vals.mean():.2f}±{vals.std():.2f} (n={len(vals)})'
        rows.append([r['变量'], ts, r[pc], r[nc], mean_sd, st, fv(r['P值'])])
    return rows, len(sig), len(cdf)

def multi_tbl(mdf, n):
    rows = []
    for _, r in mdf.iterrows():
        beta = r['β系数']; pv = r['P值']
        try:
            p_use = max(pv, 1e-30)  # 避免P=0导致Wald=inf
            wald = stats.chi2.isf(p_use, 1) if p_use < 1 else 0
        except: wald = 0
        se = abs(beta)/math.sqrt(wald) if wald > 0.01 else 0
        orv = r.get('OR', r.get('OR值', ''))
        ci = r.get('95%CI', '')
        if ci == '' or (isinstance(ci, float) and pd.isna(ci)):
            ci = f"{fv(r.get('95%CI下限',''))}-{fv(r.get('95%CI上限',''))}"
        rows.append([r['变量'], n, fv(beta), fv(se), fv(wald), fv(orv), ci, fv(pv)])
    return rows

def assign_tbl(mdf, bdf):
    rows = []
    for _, r in mdf.iterrows():
        vn = r['变量']
        m = bdf[bdf['变量']==vn]
        asg = '原始连续数值' if len(m)>0 and m.iloc[0]['类型']=='连续' else '0=否/阴性，1=是/阳性'
        rows.append([vn, asg])
    return rows

# ===== 读取数据（使用修正后的无泄漏版本） =====
bb = rc('overall_before_baseline.csv')
bu = rc('overall_before_univariate.csv')
bm = rc('overall_before_multivariate.csv')
ab = rc('overall_after_baseline.csv')
au = rc('overall_after_univariate.csv')
am = rc('overall_after_multivariate.csv')
comp = rc('compare_before_after_v2.csv')
def _fix_p(v):
    if v == '-': return v
    try:
        f = float(v)
        if f == 0 or (0 < f < 0.001): return '<0.001'
        return f'{f:.4f}'
    except: return v
comp['P值'] = comp['P值'].apply(_fix_p)
hbb = rc('havte_before_baseline.csv')
hbm = rc('havte_before_multivariate.csv')
hab = rc('havte_after_baseline.csv')
ham = rc('havte_after_multivariate.csv')
mlc = rc('ml_model_comparison.csv')
perf_orig = rc('ml_train_test_comparison.csv')
# 补充特异度列
perf_orig['特异度'] = perf_orig['TN'] / (perf_orig['TN'] + perf_orig['FP'])
perf_orig['特异度'] = perf_orig['特异度'].fillna(0)
perf_orig.rename(columns={'Accuracy':'准确率','Precision':'精确率','Recall':'灵敏度','F1':'F1值'}, inplace=True)
perf = perf_orig
shap_f = rc('shap_features.csv')
risk = rc('risk_factors_summary.csv')
ext_v = rc('external_validation_results.csv')
# 兼容处理：如果旧文件不存在则跳过
try:
    rvp = rc('r_vs_python_full_comparison.csv')
except: rvp = pd.DataFrame()
try:
    bl_comp = rc('baseline_comparison_periods.csv')
except: bl_comp = pd.DataFrame()

# 从各CSV列名动态提取n值
_pc_bb = [c for c in bb.columns if '阳性' in c][0]
_nc_bb = [c for c in bb.columns if '阴性' in c][0]
N_POS_B = int(re.search(r'n=(\d+)', _pc_bb).group(1))
N_NEG_B = int(re.search(r'n=(\d+)', _nc_bb).group(1))
N_TOT_B = N_POS_B + N_NEG_B
_pc_ab = [c for c in ab.columns if '阳性' in c][0]
_nc_ab = [c for c in ab.columns if '阴性' in c][0]
N_POS_A = int(re.search(r'n=(\d+)', _pc_ab).group(1))
N_NEG_A = int(re.search(r'n=(\d+)', _nc_ab).group(1))
N_TOT_A = N_POS_A + N_NEG_A
_pc_hbb = [c for c in hbb.columns if '阳性' in c][0]
_nc_hbb = [c for c in hbb.columns if '阴性' in c][0]
HA_B_POS = int(re.search(r'n=(\d+)', _pc_hbb).group(1))
HA_B_NEG = int(re.search(r'n=(\d+)', _nc_hbb).group(1))
HA_B_TOT = HA_B_POS + HA_B_NEG
_pc_hab = [c for c in hab.columns if '阳性' in c][0]
_nc_hab = [c for c in hab.columns if '阴性' in c][0]
HA_A_POS = int(re.search(r'n=(\d+)', _pc_hab).group(1))
HA_A_NEG = int(re.search(r'n=(\d+)', _nc_hab).group(1))
HA_A_TOT = HA_A_POS + HA_A_NEG
_train_df = pd.read_csv(os.path.join(BASE, 'train_data.csv'))
_test_df = pd.read_csv(os.path.join(BASE, 'test_data.csv'))
_train_y = _train_df['潜在可预防VTE'].map({True:1,False:0,'True':1,'False':0,1:1,0:0})
_test_y = _test_df['潜在可预防VTE'].map({True:1,False:0,'True':1,'False':0,1:1,0:0})
N_TRAIN = len(_train_df); N_TRAIN_POS = int(_train_y.sum())
N_TEST = len(_test_df); N_TEST_POS = int(_test_y.sum())

tn_ = [0]; fn_ = [0]
def nt(): tn_[0] += 1; return tn_[0]
def nf(): fn_[0] += 1; return fn_[0]

CN = {'一般':'一般资料','疾病':'疾病相关指标','治疗':'治疗相关指标','生化':'生化指标'}
SH_B = ['变量',f'合计（n={N_TOT_B}）\n阳性例数（%）',f'阳性组（n={N_POS_B}）\n阳性例数（%）',
        f'阴性组（n={N_NEG_B}）\n阳性例数（%）','均值±标准差\n（生化仅≠0，含n）','统计量值\nχ²/t','P值']
SH_A = ['变量',f'合计（n={N_TOT_A}）\n阳性例数（%）',f'阳性组（n={N_POS_A}）\n阳性例数（%）',
        f'阴性组（n={N_NEG_A}）\n阳性例数（%）','均值±标准差\n（生化仅≠0，含n）','统计量值\nχ²/t','P值']

# 读取原始数据（用于计算均值±标准差）
_orig_data = pd.read_csv(os.path.join(BASE, 'final_processed_data_full_pipeline.csv'))
_orig_data['潜在可预防VTE'] = _orig_data['潜在可预防VTE'].map({True:1, False:0, 'True':1, 'False':0})
_orig_data['入院日期_dt'] = pd.to_datetime(_orig_data['入院日期'])
# 同步数据修改（合并变量、修正等）
if '住院天数' in _orig_data.columns:
    _orig_data.loc[_orig_data['住院天数'] == -8, '住院天数'] = 8
_ca = '90天前是否我院就诊'
_cb = '本次入院前90天有无住院史、治疗史、手术史'
if _ca in _orig_data.columns and _cb in _orig_data.columns:
    _orig_data['90天内院内就诊/住院史'] = ((_orig_data[_ca]==1)|(_orig_data[_cb]==1)).astype(int)
_raw_before = _orig_data[_orig_data['入院日期_dt'] <= '2025-03-31']
_raw_after = _orig_data[_orig_data['入院日期_dt'] > '2025-03-31']
MH = ['因素','例数','β','SE','Waldχ²','OR','95%CI','P值']
MODEL_ORDER = ['Random Forest','SVM','XGBoost','Naive Bayes','Decision Tree','KNN']
MODEL_DESC = {
    'Random Forest': 'Random Forest（随机森林）是基于Bagging的集成学习方法，通过构建多棵决策树并取多数投票进行分类。参数：n_estimators=200, max_depth=10, min_samples_split=5, min_samples_leaf=2, class_weight=balanced。',
    'SVM': 'SVM（支持向量机）通过寻找最优超平面实现分类，使用RBF核函数映射到高维空间。参数：kernel=rbf, C=1.0, gamma=scale, class_weight=balanced, probability=True。',
    'XGBoost': 'XGBoost是基于梯度提升的集成学习方法，通过逐步构建弱分类器并优化损失函数。采用RandomizedSearchCV进行60组参数组合的调参优化（10折CV），最终参数由数据驱动确定。',
    'Naive Bayes': 'Naive Bayes（朴素贝叶斯）基于贝叶斯定理和特征条件独立假设。本研究采用高斯朴素贝叶斯（GaussianNB），适用于连续型特征。',
    'Decision Tree': 'Decision Tree（决策树）通过递归划分特征空间构建树结构分类器。参数：max_depth=8, min_samples_split=5, min_samples_leaf=2, class_weight=balanced。',
    'KNN': 'KNN（K近邻）基于距离度量的实例学习方法，通过最近K个邻居的多数投票分类。参数：n_neighbors=7, weights=distance, metric=minkowski(p=2)。'
}

# ====================
# 标题
# ====================
title('潜在可预防VTE影响因素分析结果')
body('')

# ====================
# 一、研究概述
# ====================
ch('一、研究概述')
body(f'本研究纳入{N_TOT_B + N_TOT_A}例VTE患者，以2025年3月31日为界划分为两个时期。3-31及之前{N_TOT_B}例（潜在可预防VTE阳性{N_POS_B}例，{N_POS_B/N_TOT_B*100:.1f}%），3-31之后{N_TOT_A}例（阳性{N_POS_A}例，{N_POS_A/N_TOT_A*100:.1f}%）。训练集与测试集均取自3-31之前数据，按80/20比例划分：训练集{N_TRAIN}例（阳性{N_TRAIN_POS}例），测试集{N_TEST}例（阳性{N_TEST_POS}例）。3-31之后{N_TOT_A}例用作外部验证集。')
body('数据预处理：（1）极端值住院天数（-8天）修正为8天，其余保留；（2）缺失的机械预防日期差值创建"_无"指示变量并填充0值；（3）删除常量列；（4）排除预防相关泄漏变量（医院相关性VTE、规范预防、是否药物/机械预防、预防措施哑变量、VTE评分与预防日期差值等71个变量），保留5个有临床意义的预防变量（规范预防、是否机械预防、是否药物预防、首次VTE中高风险评分日期与机械/药物预防日期差值）；（5）对SVM和KNN模型使用StandardScaler进行特征标准化。')

# ====================
# 二、变量筛选
# ====================
ch('二、变量筛选')

# --- 2.1 单因素 ---
sec(f'2.1 回顾性分析（n={N_TOT_B}）单因素分析（数据: overall_before_baseline.csv, overall_before_univariate.csv）')
body('注：分类变量"阳性例数"指该变量=1的人数；连续变量"均值±标准差"中生化指标仅统计检测值≠0的样本，括号内标注有效检测例数n。')
for ck in ['一般','疾病','治疗','生化']:
    si = {'一般':'2.1.1','疾病':'2.1.2','治疗':'2.1.3','生化':'2.1.4'}[ck]
    cn = CN[ck]
    sub(f'{si} {cn}')
    rows, ns, ntot = sf_table(bb, N_NEG_B, N_POS_B, ck, raw_data=_raw_before)
    t = nt()
    cap(f'表{t} {cn}单因素分析结果')
    if rows:
        mt(SH_B, rows)
        body(f'该类别共{ntot}个变量，{ns}个P<0.05。')
    else:
        body(f'该类别共{ntot}个变量，无变量P<0.05。')

# --- 2.2 多因素 ---
sec('2.2 多因素分析（数据: overall_before_multivariate.csv）')
# 【批注1】列出纳入多因素的变量
_bu_sig = bu[bu['P值'] < 0.05]
body(f'将单因素分析中P<0.05的{len(_bu_sig)}个变量纳入多因素Logistic回归。经共线性检查（r>0.8剔除）和EPV原则约束后，最终{len(bm)}个变量入模。')
# 导出纳入变量列表CSV
_enter_vars = bm[['变量']].copy()
_enter_vars['来源'] = '多因素入模'
_enter_vars.to_csv(os.path.join(BASE, 'multivariate_entered_vars.csv'), index=False, encoding='utf-8-sig')
# 在报告中列出
body('纳入多因素模型的变量：' + '、'.join(bm['变量'].tolist()) + '。')
body(f'（详见 multivariate_entered_vars.csv）')

# 2.2.1
sub('2.2.1 自变量及赋值说明')
t = nt(); cap(f'表{t} 自变量及赋值说明')
mt(['自变量','赋值说明'], assign_tbl(bm, bb))

# 2.2.2
sub('2.2.2 多因素Logistic回归分析')
sig_m = bm[bm['P值'] < 0.05]
desc = []
for _, r in sig_m.iterrows():
    _or = fv(r.get('OR', r.get('OR值', '')))
    _ci = r.get('95%CI', '')
    if _ci == '' or (isinstance(_ci, float) and pd.isna(_ci)):
        _ci = f"{fv(r.get('95%CI下限',''))}-{fv(r.get('95%CI上限',''))}"
    desc.append(f'{r["变量"]}（OR={_or}，95%CI={_ci}，P={fv(r["P值"])}）')
if desc:
    body(f'多因素分析显示（已排除预防相关泄漏变量），{len(sig_m)}个变量为独立影响因素：' + '；'.join(desc) + '。')

t = nt(); cap(f'表{t} 3-31前患者多因素Logistic回归分析')
mt(MH, multi_tbl(bm, N_TOT_B))

# 森林图
fn = nf()
img(os.path.join(BASE, 'forest_plot.png'), f'图{fn} 多因素Logistic回归森林图', 5.0)

# 2.2.3 模型构建 - 每个模型单独
sub('2.2.3 机器学习模型构建（数据: ml_train_test_comparison.csv, ml_model_comparison.csv）')
body(f'本研究选择6种经典机器学习分类模型进行风险预测。训练集{N_TRAIN}例（阳性{N_TRAIN_POS}例，阴性{N_TRAIN - N_TRAIN_POS}例），测试集{N_TEST}例（阳性{N_TEST_POS}例，阴性{N_TEST - N_TEST_POS}例）。划分依据：遵循80/20惯例（Pareto原则），确保训练集有足够数据学习特征模式，测试集提供合理评估精度。')
# 【批注2】补充K折交叉验证描述
body('为避免训练集过拟合导致性能虚高（AUC=1.0），训练集评估采用10折分层交叉验证（10-fold Stratified Cross-Validation）：将训练集随机分为10份，每次用9份训练、1份验证，循环10次取平均。该方法确保每个样本恰好被验证1次，且每折中阳性/阴性比例与整体一致（分层），所得AUC更接近模型的真实泛化能力。测试集（45例）和外部验证集（88例，3-31后数据）始终独立，未参与任何训练过程。')

for mname in MODEL_ORDER:
    body(f'【{mname}】')
    body(MODEL_DESC[mname])
    # 混淆矩阵
    fn = nf()
    fname = mname.replace(' ', '_').lower()
    img(os.path.join(BASE, f'cm_{fname}.png'), f'图{fn} {mname}测试集混淆矩阵', 3.5)

    # 解读混淆矩阵
    m_perf = perf[(perf['模型']==mname) & (perf['数据集']=='测试集')]
    if len(m_perf) > 0:
        r = m_perf.iloc[0]
        body(f'测试集：TN={int(r["TN"])}，FP={int(r["FP"])}，FN={int(r["FN"])}，TP={int(r["TP"])}。'
             f'正确识别{int(r["TP"])}例阳性和{int(r["TN"])}例阴性，'
             f'误判{int(r["FP"])}例假阳性和{int(r["FN"])}例假阴性。')

    # 模型评价指标表（训练+测试）
    t = nt()
    cap(f'表{t} {mname}模型评价指标')
    m_rows = []
    for _, pr in perf[perf['模型']==mname].iterrows():
        f1v = pr.get('F1', pr.get('F1值', 0))
        m_rows.append([pr['数据集'], fv(pr['准确率']), fv(pr['精确率']),
                       fv(pr['灵敏度']), fv(pr['特异度']), fv(f1v), fv(pr['AUC'])])
    mt(['','准确率','精确率','灵敏度','特异度','F1','AUC'], m_rows)

    # ROC曲线
    fn = nf()
    img(os.path.join(BASE, f'roc_{fname}.png'), f'图{fn} {mname} ROC曲线（训练+测试）', 4.0)

# 2.2.4 模型评价
sub('2.2.4 模型评价')
# 动态生成模型排序文本
_test_perf = perf[perf['数据集']=='测试集'].sort_values('AUC', ascending=False)
_perf_strs = [f'{r["模型"]}（AUC={r["AUC"]:.3f}）' for _, r in _test_perf.iterrows()]
_best = _test_perf.iloc[0]
body(f'综合测试集评估指标（训练集采用10折交叉验证），模型性能排序为：{"  > ".join(_perf_strs)}。{_best["模型"]}综合判别性能最优，F1={_best.get("F1值",_best.get("F1",0)):.3f}，灵敏度={_best["灵敏度"]:.3f}。')

t = nt()
cap(f'表{t} 各模型测试集性能指标对比')
test_perf = perf[perf['数据集']=='测试集'].sort_values('AUC', ascending=False)
cr = []
for _, r in test_perf.iterrows():
    f1v = r.get('F1', r.get('F1值', 0))
    cr.append([r['模型'], fv(r['准确率']), fv(r['精确率']), fv(r['灵敏度']),
               fv(r['特异度']), fv(f1v), fv(r['AUC'])])
mt(['分类模型','准确率','精确率','灵敏度','特异度','F1','AUC'], cr)

# 2.2.5 SHAP可解释性
sub('2.2.5 预测模型的可解释性（数据: shap_features.csv, risk_factors_summary.csv）')
body('使用SHAP（SHapley Additive exPlanations）对最佳模型XGBoost进行可解释性分析。SHAP基于博弈论中的Shapley值，为每个特征的每个样本计算其对模型预测的边际贡献，具有局部一致性和全局可加性。')
# 动态生成SHAP Top5文本
_shap_top5 = shap_f.head(5)
_shap_strs = []
for i, (_, sr_) in enumerate(_shap_top5.iterrows(), 1):
    _shap_strs.append(f'（{i}）{sr_["变量"]}（均值={sr_["SHAP均值"]:.3f}）')
body(f'SHAP分析Top5特征（已排除预防相关泄漏变量）：{"；".join(_shap_strs)}。各特征均为真实临床变量，与多因素Logistic回归结论一致。')

fn = nf()
img(os.path.join(BASE, 'shap_importance.png'), f'图{fn} XGBoost SHAP特征重要性Top20', 5.0)
fn = nf()
img(os.path.join(BASE, 'shap_summary_v2.png'), f'图{fn} SHAP概要图（蜂群图）', 5.5)
fn = nf()
img(os.path.join(BASE, 'shap_interaction.png'), f'图{fn} SHAP交互效应热力图', 5.0)

# 【批注3】单因素+多因素+SHAP有意义汇总表（csv+报告表格）
sec('2.3 3-31前分析结果汇总（数据: key_factors_summary.csv）')
body('将单因素（P<0.05）、多因素（入模变量）和SHAP重要性Top20进行综合汇总：')

# 构建汇总数据
_uni_sig_b = bu[bu['P值'] < 0.05][['变量', 'P值']].copy()
_uni_sig_b.columns = ['变量', '单因素P值']
_or_col = 'OR值' if 'OR值' in bu.columns else 'OR'
_uni_sig_b['单因素OR'] = bu[bu['P值'] < 0.05][_or_col].values

_or_col_m = 'OR值' if 'OR值' in bm.columns else 'OR'
_multi_b = bm[['变量', 'P值', _or_col_m]].copy()
_multi_b.columns = ['变量', '多因素P值', '多因素OR']

_shap_top = shap_f.head(20)[['变量', 'SHAP均值']].copy()

# 合并
_summary_all = _shap_top.merge(_uni_sig_b, on='变量', how='outer')
_summary_all = _summary_all.merge(_multi_b, on='变量', how='outer')
_summary_all = _summary_all.sort_values('SHAP均值', ascending=False, na_position='last')
_summary_all.to_csv(os.path.join(BASE, 'key_factors_summary.csv'), index=False, encoding='utf-8-sig')

# 报告表格
t = nt()
cap(f'表{t} 3-31前关键影响因素综合汇总（单因素+多因素+SHAP）')
_sum_rows = []
for _, r in _summary_all.head(20).iterrows():
    _sum_rows.append([
        r['变量'],
        fv(r.get('SHAP均值', '')),
        fv(r.get('单因素OR', '')),
        fv(r.get('单因素P值', '')),
        fv(r.get('多因素OR', '')),
        fv(r.get('多因素P值', ''))
    ])
mt(['变量', 'SHAP均值', '单因素OR', '单因素P值', '多因素OR', '多因素P值'], _sum_rows)
body(f'（详见 key_factors_summary.csv）')

# ====================
# 三、3-31后分析结果
# ====================
ch(f'三、3-31后分析结果（n={N_TOT_A}）')

# 3.1 单因素
sec('3.1 特征分析（数据: overall_after_baseline.csv, overall_after_univariate.csv）')
for ck in ['一般','疾病','治疗','生化']:
    si = {'一般':'3.1.1','疾病':'3.1.2','治疗':'3.1.3','生化':'3.1.4'}[ck]
    cn = CN[ck]
    sub(f'{si} {cn}')
    rows, ns, ntot = sf_table(ab, N_NEG_A, N_POS_A, ck, raw_data=_raw_after)
    t = nt()
    cap(f'表{t} {cn}单因素分析结果')
    if rows:
        mt(SH_A, rows); body(f'该类别共{ntot}个变量，{ns}个P<0.05。')
    else:
        body(f'该类别共{ntot}个变量，无变量P<0.05。')

# 3.2 多因素
sec('3.2 多因素分析（数据: overall_after_multivariate.csv）')
body('将3-31后单因素分析中P<0.05的变量纳入多因素Logistic回归。')

sub('3.2.1 自变量及赋值说明')
t = nt(); cap(f'表{t} 自变量及赋值说明')
mt(['自变量','赋值说明'], assign_tbl(am, ab))

sub('3.2.2 多因素Logistic回归分析')
sig_a = am[am['P值'] < 0.05]
if len(sig_a) > 0:
    d2 = []
    for _, r in sig_a.iterrows():
        _or = fv(r.get('OR', r.get('OR值', '')))
        _ci = r.get('95%CI', '')
        if _ci == '' or (isinstance(_ci, float) and pd.isna(_ci)):
            _ci = f"{fv(r.get('95%CI下限',''))}-{fv(r.get('95%CI上限',''))}"
        d2.append(f'{r["变量"]}（OR={_or}，95%CI={_ci}，P={fv(r["P值"])}）')
    body(f'多因素分析显示，{len(sig_a)}个变量达到统计学显著：' + '；'.join(d2) + '。')
else:
    body('多因素分析中无变量达到P<0.05，趋势性显著：')
    for _, r in am.head(3).iterrows():
        _or = fv(r.get('OR', r.get('OR值', '')))
        body(f'  {r["变量"]}（OR={_or}，P={fv(r["P值"])}）')

t = nt(); cap(f'表{t} 3-31后患者多因素Logistic回归分析')
mt(MH, multi_tbl(am, N_TOT_A))

# 3.3 结局指标分析
sec('3.3 结局指标分析（数据: compare_before_after_v2.csv, havte_before/after_*.csv）')

# 3.3.1 前后基线对比
sub('3.3.1 3-31前后基线对比')
if not bl_comp.empty:
    body('对3-31前后两个时期的关键基线变量进行对比：')
    t = nt(); cap(f'表{t} 3-31前后基线对比')
    bl_rows = [[str(v) for v in r] for _, r in bl_comp.iterrows()]
    mt(list(bl_comp.columns), bl_rows)
else:
    body('（基线对比数据暂未生成）')

# 3.3.2 HA-VTE
sub('3.3.2 HA-VTE发生率前后对比')
t = nt(); cap(f'表{t} HA-VTE 3-31前后对比')
hv_r = comp[comp['指标'].str.contains('HA-VTE|总VTE')]
mt(list(comp.columns), [[str(v) for v in r] for _, r in hv_r.iterrows()])

# 3.3.3 潜在可预防VTE
sub('3.3.3 潜在可预防VTE发生率前后对比')
t = nt(); cap(f'表{t} 潜在可预防VTE 3-31前后对比')
pv_r = comp[comp['指标'].str.contains('潜在可预防|规范预防')]
mt(list(comp.columns), [[str(v) for v in r] for _, r in pv_r.iterrows()])
# 动态生成潜在可预防VTE前后对比描述
_comp_ppv3 = comp[comp['指标'].str.contains('潜在可预防VTE在HA-VTE')]
_comp_norm3 = comp[comp['指标'].str.contains('潜在可预防VTE阳性中规范')]
_desc_parts = []
if len(_comp_ppv3) > 0:
    _rp = _comp_ppv3.iloc[0]
    _desc_parts.append(f'潜在可预防VTE在HA-VTE中占比：3-31前{_rp["3-31前"]}，3-31后{_rp["3-31后"]}（P={_rp["P值"]}）')
if len(_comp_norm3) > 0:
    _rn3 = _comp_norm3.iloc[0]
    _desc_parts.append(f'规范预防率：3-31前{_rn3["3-31前"]}，3-31后{_rn3["3-31后"]}（P={_rn3["P值"]}）')
if _desc_parts:
    body('。'.join(_desc_parts) + '。')

# 3.3.4 规范预防（HA-VTE=1）
sub('3.3.4 规范预防影响因素分析（HA-VTE=1，目标变量：规范预防）')
body(f'对HA-VTE=1患者进行规范预防影响因素分析。3-31前{HA_B_TOT}例（规范预防{HA_B_POS}例），3-31后{HA_A_TOT}例（规范预防{HA_A_POS}例）。')

# 【批注4】3-31前 - 增加均值±标准差列，生化例数只填≠0的
body(f'（一）3-31前（n={HA_B_TOT}）')
hb_sig = hbb[hbb['P值'] < 0.05]
nc_hb = [c for c in hbb.columns if '阴性' in c][0]
pc_hb = [c for c in hbb.columns if '阳性' in c][0]

# 读取原始数据计算均值±标准差
_orig = pd.read_csv(os.path.join(BASE, 'final_processed_data_full_pipeline.csv'))
_orig['潜在可预防VTE'] = _orig['潜在可预防VTE'].map({True:1, False:0, 'True':1, 'False':0})
_orig['入院日期_dt'] = pd.to_datetime(_orig['入院日期'])
_havte_b = _orig[(_orig['入院日期_dt'] <= '2025-03-31') & ((_orig['医院相关性VTE'] == 1) | (_orig['我院相关VTE'] == 1))].copy()

# 生化关键词（用于判断是否只填≠0的例数）
_bio_kw = ['嗜酸','D-二聚体','DD','APTT','凝血','钾离子','纤维蛋白原','白蛋白','血红蛋白',
           '总钙','白细胞','中性粒','红细胞','Fbg','FEU','PT','MCH','MCHC','Braden']

def _is_bio(vname):
    return any(k in vname for k in _bio_kw)

hb_rows = []
for _, r in hb_sig.iterrows():
    ts, st = total_stat(r, HA_B_NEG, HA_B_POS)
    vname = r['变量']
    # 计算均值±标准差
    if r['类型'] == '连续' and vname in _havte_b.columns:
        vals = _havte_b[vname].dropna()
        if _is_bio(vname):
            vals = vals[vals != 0]  # 生化只取≠0
        mean_sd = f'{vals.mean():.2f}±{vals.std():.2f}' if len(vals) > 0 else '-'
        n_valid = len(vals)
    elif r['类型'] == '分类' and vname in _havte_b.columns:
        mean_sd = '-'
        n_valid = int(_havte_b[vname].sum())
    else:
        mean_sd = '-'
        n_valid = '-'
    hb_rows.append([r['变量'], ts, r[pc_hb], r[nc_hb], mean_sd, st, fv(r['P值'])])

t = nt(); cap(f'表{t} 3-31前HA-VTE规范预防单因素分析（P<0.05）')
mt(['变量',f'合计（n={HA_B_TOT}）\n例数（%）',f'规范预防组（n={HA_B_POS}）\n例数（%）',
    f'非规范预防组（n={HA_B_NEG}）\n例数（%）','均值±标准差\n（生化仅≠0）','统计量值\nχ²/t','P值'], hb_rows)
body(f'共{len(hbb)}个变量，{len(hb_sig)}个P<0.05。注：生化指标均值±标准差仅统计检测值≠0的样本。')

t = nt(); cap(f'表{t} 3-31前HA-VTE规范预防多因素Logistic回归')
mt(MH, multi_tbl(hbm, HA_B_TOT))
for _, r in hbm[hbm['P值']<0.05].iterrows():
    _or = fv(r.get('OR', r.get('OR值', '')))
    _ci = r.get('95%CI', '')
    if _ci == '' or (isinstance(_ci, float) and pd.isna(_ci)):
        _ci = f"{fv(r.get('95%CI下限',''))}-{fv(r.get('95%CI上限',''))}"
    body(f'  {r["变量"]}：OR={_or}，95%CI={_ci}，P={fv(r["P值"])}')

# 【批注5】3-31后 - 同表23增加均值±标准差列
body(f'（二）3-31后（n={HA_A_TOT}）')
body('样本量较小，结果需谨慎解读。')
ha_sig = hab[hab['P值'] < 0.05]
nc_ha = [c for c in hab.columns if '阴性' in c][0]
pc_ha = [c for c in hab.columns if '阳性' in c][0]
ha_nn = int(re.search(r'n=(\d+)', nc_ha).group(1))
ha_np = int(re.search(r'n=(\d+)', pc_ha).group(1))

# 3-31后HA-VTE原始数据
_havte_a = _orig[(_orig['入院日期_dt'] > '2025-03-31') & ((_orig['医院相关性VTE'] == 1) | (_orig['我院相关VTE'] == 1))].copy()

ha_rows = []
for _, r in ha_sig.iterrows():
    ts, st = total_stat(r, ha_nn, ha_np)
    vname = r['变量']
    if r['类型'] == '连续' and vname in _havte_a.columns:
        vals = _havte_a[vname].dropna()
        if _is_bio(vname):
            vals = vals[vals != 0]
        mean_sd = f'{vals.mean():.2f}±{vals.std():.2f}' if len(vals) > 0 else '-'
    else:
        mean_sd = '-'
    ha_rows.append([r['变量'], ts, r[pc_ha], r[nc_ha], mean_sd, st, fv(r['P值'])])
t = nt(); cap(f'表{t} 3-31后HA-VTE规范预防单因素分析（P<0.05）')
if ha_rows:
    mt(['变量',f'合计（n={HA_A_TOT}）\n例数（%）',f'规范预防组（n={HA_A_POS}）\n例数（%）',
        f'非规范预防组（n={HA_A_NEG}）\n例数（%）','均值±标准差\n（生化仅≠0）','统计量值\nχ²/t','P值'], ha_rows)
else:
    body('无变量达到P<0.05。')
body(f'共{len(hab)}个变量，{len(ha_sig)}个P<0.05。')

t = nt(); cap(f'表{t} 3-31后HA-VTE规范预防多因素Logistic回归')
mt(MH, multi_tbl(ham, HA_A_TOT))
sig_ham = ham[ham['P值']<0.05]
if len(sig_ham) > 0:
    for _, r in sig_ham.iterrows():
        _or = fv(r.get('OR', r.get('OR值', '')))
        body(f'  {r["变量"]}：OR={_or}，P={fv(r["P值"])}')
else:
    body('无变量达到统计学显著（P<0.05），可能与样本量过小有关。')

# ====================
# 四、结论
# ====================
ch('四、结论')
# 动态生成结论
# 3-31前多因素
_bm_sig = bm[bm['P值'] < 0.05]
body(f'（一）潜在可预防VTE独立影响因素（3-31前{N_TOT_B}例）：')
for _, _r in _bm_sig.iterrows():
    _or = _r.get('OR', _r.get('OR值', ''))
    _ci_lo = _r.get('95%CI下限', '')
    _ci_hi = _r.get('95%CI上限', '')
    _role = '独立危险因素' if float(_or) > 1 else '独立保护因素'
    body(f'  • {_role}：{_r["变量"]}（OR={fv(_or)}，95%CI={fv(_ci_lo)}-{fv(_ci_hi)}，P={fv(_r["P值"])}）')

# 3-31后多因素
_am_sig = am[am['P值'] < 0.05]
if len(_am_sig) > 0:
    _strs = []
    for _, _r in _am_sig.iterrows():
        _or = _r.get('OR', _r.get('OR值', ''))
        _strs.append(f'{_r["变量"]}（OR={fv(_or)}，P={fv(_r["P值"])}）')
    body(f'（二）3-31后{N_TOT_A}例：{"\u3001".join(_strs)}。')
else:
    body(f'（二）3-31后{N_TOT_A}例：无变量达到P<0.05。')

# HA-VTE规范预防
_hbm_sig = hbm[hbm['P值'] < 0.05]
if len(_hbm_sig) > 0:
    _strs = []
    for _, _r in _hbm_sig.iterrows():
        _or = _r.get('OR', _r.get('OR值', ''))
        _strs.append(f'{_r["变量"]}（OR={fv(_or)}，P={fv(_r["P值"])}）')
    body(f'（三）规范预防影响因素（3-31前HA-VTE {HA_B_TOT}例）：{"\u3001".join(_strs)}。')

# 动态生成前后对比结论
_comp_ppv = comp[comp['指标'].str.contains('潜在可预防VTE在HA-VTE')]
_comp_norm = comp[comp['指标'].str.contains('潜在可预防VTE阳性中规范')]
if len(_comp_ppv) > 0:
    _r = _comp_ppv.iloc[0]
    body(f'（四）前后对比：潜在可预防VTE在HA-VTE中占比，3-31前{_r["3-31前"]}，3-31后{_r["3-31后"]}（P={_r["P值"]}）。')
else:
    body('（四）前后对比数据请参见表格。')
if len(_comp_norm) > 0:
    _rn = _comp_norm.iloc[0]
    body(f'     规范预防率：3-31前{_rn["3-31前"]}，3-31后{_rn["3-31后"]}（P={_rn["P值"]}）。')

# ML结论 - 动态
_best_test = mlc.sort_values('AUC', ascending=False).iloc[0]
_best_ext = ext_v.sort_values('AUC', ascending=False).iloc[0]
body(f'（五）机器学习（训练集10折交叉验证）：{_best_test["模型"]}最佳（测试集AUC={_best_test["AUC"]:.3f}），SHAP与统计分析结果高度一致。外部验证最佳AUC={_best_ext["AUC"]:.3f}（{_best_ext["模型"]}）。')
body('（六）临床建议：加强90天内就诊患者VTE预防；规范开展出血风险评估；提高规范预防率。')

# ====================
# 五、方法一致性验证
# ====================
ch('五、方法一致性验证')
# 动态统计单因素显著变量数
_bu_sig = len(bu[bu['P值'] < 0.05]) if not bu.empty else 0
_au_sig = len(au[au['P值'] < 0.05]) if not au.empty else 0
_hbu = rc('havte_before_univariate.csv')
_hau = rc('havte_after_univariate.csv')
_hbu_sig = len(_hbu[_hbu['P值'] < 0.05]) if not _hbu.empty else 0
_hau_sig = len(_hau[_hau['P值'] < 0.05]) if not _hau.empty else 0

# 【批注6】泄漏变量列表CSV+报告表格
_leak_kw = ['预防', '医院相关性VTE', '我院相关VTE']
_keep_vars = ['规范预防', '是否机械预防', '是否药物预防',
    '首次VTE中高风险评分日期与机械预防日期差值',
    '首次VTE中高风险评分日期与药物预防日期差值']
_all_cols_orig = pd.read_csv(os.path.join(BASE, 'final_processed_data_full_pipeline.csv'), nrows=0).columns.tolist()
_leaked = []
for c in _all_cols_orig:
    if c in _keep_vars:
        continue
    for kw in _leak_kw:
        if kw in c:
            _reason = '定义前提（阳性组100%=1）' if '医院相关性' in c or '我院相关' in c else '预防相关变量（直接/间接参与目标定义）'
            _leaked.append({'排除变量': c, '排除原因': _reason})
            break
_leak_df = pd.DataFrame(_leaked)
_leak_df.to_csv(os.path.join(BASE, 'excluded_leak_variables.csv'), index=False, encoding='utf-8-sig')

body(f'本分析已排除{len(_leak_df)}个预防相关泄漏变量，保留5个有临床意义的预防变量（规范预防、是否机械预防、是否药物预防、评分与预防日期差值）。')
body(f'（详见 excluded_leak_variables.csv）')

# 在报告中列出关键泄漏变量
t = nt(); cap(f'表{t} 排除的泄漏变量列表（部分）')
_leak_show = _leak_df.head(15)
_leak_rows = [[r['排除变量'], r['排除原因']] for _, r in _leak_show.iterrows()]
_leak_rows.append([f'...共{len(_leak_df)}个', ''])
mt(['排除变量', '排除原因'], _leak_rows)
body(f'（1）3-31前{N_TOT_B}例：单因素{_bu_sig}个变量P<0.05。')
body(f'（2）3-31后{N_TOT_A}例：单因素{_au_sig}个变量P<0.05。')
body(f'（3）HA-VTE规范预防（3-31前{HA_B_TOT}例）：单因素{_hbu_sig}个变量P<0.05。')
body(f'（4）HA-VTE规范预防（3-31后{HA_A_TOT}例）：单因素{_hau_sig}个变量P<0.05。')
body('单因素分析：分类变量采用χ²检验（期望频数<5时改用Fisher精确检验），连续变量采用独立样本t检验（正态分布）或Mann-Whitney U检验（非正态分布）。')
body('多因素分析：采用二元Logistic回归，共线性检查（r>0.8剔除），EPV原则约束入模变量数。')
body(f'数据预处理：（1）极端值住院天数（-8天）修正为8天；（2）二分类变量0/1编码；（3）多分类变量One-Hot编码；（4）缺失的机械预防日期差值创建\"_无\"指示变量；（5）删除常量列；（6）排除{len(_leak_df)}个预防相关泄漏变量，保留5个有临床意义的预防变量。')

if not rvp.empty:
    body('Python与R机器学习交叉验证结果：')
    t = nt(); cap(f'表{t} Python与R语言6种模型AUC全面对比')
    mt(list(rvp.columns), [[str(v) for v in r] for _, r in rvp.iterrows()])

# ====================
# 保存
# ====================
out = os.path.join(BASE, 'VTE影响因素分析报告.docx')
doc.save(out)
print(f'报告已生成: {out}')
print(f'表格: {tn_[0]}, 图片: {fn_[0]}')

doc2 = Document(out)
rt = len(doc2.tables)
ri = sum(1 for r in doc2.part.rels.values() if 'image' in r.reltype)
rp = len([p for p in doc2.paragraphs if p.text.strip()])
print(f'实际: 表格={rt}, 图片={ri}, 段落={rp}')
