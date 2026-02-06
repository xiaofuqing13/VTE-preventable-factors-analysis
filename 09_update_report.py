# -*- coding: utf-8 -*-
"""
09_update_report.py
在现有Word报告末尾追加机器学习分析章节
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'C:\Users\Administrator\Desktop\2.5.388'

def set_cell_font(cell, text, size=10, bold=False, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_heading_black(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table_from_df(doc, df, caption=None):
    if caption:
        add_para(doc, caption, size=10, bold=True)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        set_cell_font(table.rows[0].cells[j], col, size=9, bold=True)
    for i in range(len(df)):
        for j, col in enumerate(df.columns):
            val = df.iloc[i, j]
            if isinstance(val, float):
                val = f'{val:.4f}' if abs(val) < 10 else f'{val:.2f}'
            set_cell_font(table.rows[i + 1].cells[j], str(val), size=9)
    return table

def add_image(doc, path, caption=None, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            add_para(doc, caption, size=9, bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

# 打开现有报告
doc = Document(os.path.join(BASE, 'VTE影响因素分析报告.docx'))

# ============================================================
# 第十章 机器学习模型构建与评估
# ============================================================
add_heading_black(doc, '第十章 机器学习预测模型构建与评估', level=1)

# 10.1 数据划分与特征工程
add_heading_black(doc, '10.1 数据划分与特征工程', level=2)
add_para(doc, '本研究使用3月31日前收集的224例数据构建机器学习预测模型，目标变量为"潜在可预防VTE"。数据集按80/20比例划分为训练集（179例，其中阳性82例，阴性97例）和测试集（45例，其中阳性20例，阴性25例）。')
add_para(doc, '数据划分比例选择依据：（1）遵循机器学习领域广泛采用的80/20划分惯例（Pareto原则）；（2）考虑到样本量较小（n=224），80%训练集可确保模型有足够数据学习特征模式；（3）20%测试集（n=45）在本研究正负样本比例下可提供合理的评估精度。')
add_para(doc, '特征工程处理：（1）删除非数值列（入院日期、预防措施、dataset）；（2）对缺失的机械预防日期差值创建"_无"指示变量并填充0值；（3）删除常量列；（4）对训练集、测试集和外部验证集取公共特征列，最终保留414个特征变量。对SVM和KNN模型使用StandardScaler进行特征标准化。')

# 10.2 模型选择与参数
add_heading_black(doc, '10.2 模型选择与参数设置', level=2)
add_para(doc, '本研究选择6种经典机器学习分类模型进行潜在可预防VTE的风险预测，涵盖集成学习、核方法、概率模型、实例学习等不同算法范式，以全面评估不同方法的预测性能。')

params = pd.read_csv(os.path.join(BASE, 'ml_model_params.csv'))
add_table_from_df(doc, params, caption='表21 机器学习模型参数设置')

add_para(doc, '模型参数设置说明：（1）Random Forest和Decision Tree使用class_weight=balanced自动平衡类别权重；（2）XGBoost通过scale_pos_weight=1.18调整正负样本权重；（3）SVM使用RBF核函数并启用概率估计；（4）Naive Bayes选用高斯分布假设，适用于连续型特征；（5）KNN使用距离加权的7近邻策略。')

# 10.3 模型性能评估
add_heading_black(doc, '10.3 模型训练集与测试集性能对比', level=2)

tt = pd.read_csv(os.path.join(BASE, 'ml_train_test_comparison.csv'))
add_table_from_df(doc, tt, caption='表22 各模型训练集与测试集性能对比')

add_para(doc, '训练集与测试集性能分析：')
add_para(doc, '（1）XGBoost表现最佳：测试集AUC=0.982，F1=0.884，Accuracy=0.889，Recall=0.95，具有最优的综合判别能力，尤其在敏感性方面表现突出，仅漏诊1例。')
add_para(doc, '（2）Random Forest和Decision Tree测试集AUC分别为0.928和0.845，表现良好。Random Forest精确率最高（0.875），Decision Tree在敏感性（0.85）和精确率（0.810）间取得较好平衡。')
add_para(doc, '（3）SVM测试集AUC=0.802，性能中等。Naive Bayes测试集AUC=0.675，判别能力较弱，假阳性率较高（FP=10）。')
add_para(doc, '（4）KNN测试集AUC=0.518，接近随机水平，F1=0.000，未能正确预测任何阳性样本，考虑与高维稀疏特征空间中距离度量失效（"维度灾难"）有关。')
add_para(doc, '（5）多数模型训练集AUC=1.000，存在一定过拟合，但XGBoost测试集性能仍然优异（AUC=0.982），说明其泛化能力较强。')

# 10.4 模型间测试集对比
add_heading_black(doc, '10.4 模型间性能对比（测试集）', level=2)

comp = pd.read_csv(os.path.join(BASE, 'ml_model_comparison.csv'))
add_table_from_df(doc, comp, caption='表23 六种模型测试集性能对比（按AUC排序）')

add_para(doc, '综合测试集评估指标，模型性能排序为：XGBoost > Random Forest > Decision Tree > SVM > Naive Bayes > KNN。XGBoost以AUC=0.982显著优于其他模型，被选为最佳模型用于后续SHAP可解释性分析和外部验证。')

add_image(doc, os.path.join(BASE, 'roc_comparison.png'), caption='图1 测试集ROC曲线对比', width=5)
add_image(doc, os.path.join(BASE, 'roc_train.png'), caption='图2 训练集ROC曲线对比', width=5)
add_image(doc, os.path.join(BASE, 'confusion_matrices.png'), caption='图3 六种模型测试集混淆矩阵', width=6)
add_image(doc, os.path.join(BASE, 'train_test_comparison.png'), caption='图4 训练集与测试集AUC/F1对比柱状图', width=5.5)

# 10.5 SHAP可解释性
add_heading_black(doc, '10.5 SHAP可解释性分析', level=2)
add_para(doc, '使用SHAP（SHapley Additive exPlanations）方法对最佳模型XGBoost进行可解释性分析。SHAP基于博弈论中的Shapley值，能够为每个特征的每个样本计算其对模型预测的边际贡献，具有局部一致性和全局可加性的优良数学性质。')

shap = pd.read_csv(os.path.join(BASE, 'shap_features.csv'))
shap_top15 = shap.head(15)
add_table_from_df(doc, shap_top15, caption='表24 XGBoost模型SHAP特征重要性Top15')

add_para(doc, 'SHAP分析结果显示，对XGBoost模型预测贡献最大的前5个特征为：')
add_para(doc, '（1）90天前是否我院就诊（SHAP均值=1.951）：对模型预测影响最大，与单因素（OR=20.12）和多因素分析（OR=38.24）结果高度一致，提示近期就诊史是潜在可预防VTE最强的危险因素。')
add_para(doc, '（2）是否机械预防（SHAP均值=1.045）：保护性因素，与多因素分析OR=0.22一致，说明机械预防措施能有效降低潜在可预防VTE风险。')
add_para(doc, '（3）医院相关性VTE（SHAP均值=1.028）：在逻辑回归模型中因共线性未纳入，但机器学习模型能识别其独立贡献。')
add_para(doc, '（4）出血风险评估（SHAP均值=0.662）：与多因素分析OR=0.30一致，进行出血风险评估可降低VTE风险。')
add_para(doc, '（5）VTE诊断日期与入院日期是否大于24h（SHAP均值=0.642）：VTE诊断时间特征对预测有重要贡献。')

add_image(doc, os.path.join(BASE, 'shap_summary.png'), caption='图5 XGBoost模型SHAP Summary Plot（Top20）', width=5.5)

# 10.6 危险因素汇总
add_heading_black(doc, '10.6 危险因素综合汇总（单因素+多因素+SHAP）', level=2)
add_para(doc, '将SHAP特征重要性Top20与单因素、多因素Logistic回归显著变量进行综合对比，以验证机器学习模型发现的重要因素与传统统计学方法的一致性。')

risk = pd.read_csv(os.path.join(BASE, 'risk_factors_summary.csv'))
add_table_from_df(doc, risk, caption='表25 危险因素综合汇总表（SHAP Top20 + 单多因素OR）')

add_para(doc, '综合汇总分析发现：')
add_para(doc, '（1）90天前是否我院就诊在三种方法中均排名第一（SHAP=1.95, 单因素OR=20.12, 多因素OR=38.24），是最稳健的危险因素。')
add_para(doc, '（2）是否机械预防和出血风险评估在三种方法中均显著，验证了保护性因素的一致性。')
add_para(doc, '（3）SHAP识别出部分传统分析未发现的重要变量（如身高、APTT、总钙），可能反映了非线性交互效应。')
add_para(doc, '（4）总体而言，机器学习方法与传统统计方法的危险因素识别结果具有高度一致性，增强了结论的可靠性。')

add_image(doc, os.path.join(BASE, 'forest_plot.png'), caption='图6 多因素Logistic回归森林图', width=5)

# 10.7 外部验证
add_heading_black(doc, '10.7 外部验证（3月31日后88例）', level=2)
add_para(doc, '使用3月31日后收集的88例数据（阳性63例，阴性25例）作为独立外部验证集，评估各模型的泛化能力和时间迁移性能。')

ext = pd.read_csv(os.path.join(BASE, 'external_validation_results.csv'))
add_table_from_df(doc, ext, caption='表26 六种模型外部验证性能')

add_para(doc, '外部验证结果分析：')
add_para(doc, '（1）Random Forest在外部验证中表现最佳（AUC=0.743），其次为SVM（AUC=0.707）和XGBoost（AUC=0.691），三者差异不大。')
add_para(doc, '（2）所有模型在外部验证中性能均低于测试集，提示3月31日前后的患者特征分布存在一定差异（时间漂移效应）。')
add_para(doc, '（3）各模型精确率普遍较高（RF=0.971, SVM=0.938, XGBoost=0.912），说明预测为阳性的病例可靠性高，但敏感性较低（RF=0.540, XGBoost=0.492），存在漏诊风险。')
add_para(doc, '（4）外部验证数据正负样本比例（63:25≈2.5:1）与训练数据（82:97≈0.85:1）差异较大，可能是性能下降的主要原因。')

add_image(doc, os.path.join(BASE, 'external_roc.png'), caption='图7 外部验证ROC曲线对比', width=5)

# 10.8 R语言交叉验证
add_heading_black(doc, '10.8 R语言交叉验证', level=2)
add_para(doc, '使用R语言（R 4.5.2）独立实现Random Forest和XGBoost模型，验证Python结果的可重复性。R语言使用randomForest包和xgboost包，采用相同的参数设置。')

r_val = pd.read_csv(os.path.join(BASE, 'r_validation_comparison.csv'))
add_table_from_df(doc, r_val, caption='表27 Python与R语言模型性能对比')

add_para(doc, 'Python与R验证结果显示，两种语言实现的模型AUC差异均<0.05，XGBoost测试集AUC完全一致（均为0.982），Random Forest测试集AUC差异仅为0.016，验证了分析结果的可靠性和可重复性。微小差异源于两种语言的随机数生成器和算法实现细节不同。')

# 10.9 小结
add_heading_black(doc, '10.9 小结', level=2)
add_para(doc, '本章使用6种机器学习模型对潜在可预防VTE进行风险预测分析，主要发现：')
add_para(doc, '（1）XGBoost为最佳预测模型：测试集AUC=0.982，F1=0.884，综合判别性能显著优于其他模型。')
add_para(doc, '（2）SHAP可解释性分析揭示了关键预测特征，与传统单因素/多因素分析结果高度一致：90天前就诊史、机械预防、出血风险评估是最重要的因素。')
add_para(doc, '（3）外部验证（3-31后88例）显示模型具有一定泛化能力（最佳AUC=0.743），但受样本分布变化影响性能有所下降。')
add_para(doc, '（4）R语言交叉验证证实了Python分析结果的可重复性。')
add_para(doc, '（5）综合统计学方法和机器学习方法，90天前就诊史、机械预防、出血风险评估是潜在可预防VTE最稳健的影响因素，建议临床重点关注。')

# 保存
output_path = os.path.join(BASE, 'VTE影响因素分析报告.docx')
doc.save(output_path)
print(f'报告已更新保存至: {output_path}')
print('新增内容: 第十章 机器学习预测模型构建与评估')
print('  10.1 数据划分与特征工程')
print('  10.2 模型选择与参数设置 (表21)')
print('  10.3 训练集与测试集性能对比 (表22)')
print('  10.4 模型间性能对比 (表23, 图1-4)')
print('  10.5 SHAP可解释性分析 (表24, 图5)')
print('  10.6 危险因素综合汇总 (表25, 图6)')
print('  10.7 外部验证 (表26, 图7)')
print('  10.8 R语言交叉验证 (表27)')
print('  10.9 小结')
